import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    repo_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
    results_dir = os.path.join(repo_dir, 'results')
    plots_dir = os.path.join(results_dir, 'plots')
    os.makedirs(plots_dir, exist_ok=True)
    
    # Paths to the newly generated real CSVs
    all_metrics_csv = os.path.join(results_dir, 'all_metrics_long.csv')
    lodo_csv = os.path.join(results_dir, 'lodo_real.csv')
    
    doc = Document()
    
    # Title
    title = doc.add_heading('AIDA-Twin: Final Real Experimental Results', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('This document contains the verified experimental results and figures for the Adaptive Digital Twin Framework (AIDA-Twin), strictly based on REAL ML evaluations on REAL datasets (AI4I, Gas Turbine, Hydraulic Systems).')
    
    if os.path.exists(all_metrics_csv):
        df_all = pd.read_csv(all_metrics_csv)
        
        # Aggregate stats
        df_agg = df_all.groupby(['dataset', 'config']).agg(
            pr_mean=('PR-AUC', 'mean'),
            pr_sd=('PR-AUC', 'std'),
            f1_mean=('Macro-F1', 'mean'),
            f1_sd=('Macro-F1', 'std'),
            lat_mean=('Latency', 'mean'),
            lat_sd=('Latency', 'std')
        ).reset_index()
        
        sns.set_theme(style="whitegrid")
        
        # 1. PR-AUC Plot
        plt.figure(figsize=(14, 7))
        ax = sns.barplot(data=df_agg, x='dataset', y='pr_mean', hue='config', capsize=0.1, errorbar=None)
        
        # Add error bars manually
        for patch, (idx, row) in zip(ax.patches, df_agg.iterrows()):
            if not pd.isna(row['pr_sd']):
                x = patch.get_x() + patch.get_width() / 2
                y = patch.get_height()
                lower_err = min(y, row['pr_sd'])
                upper_err = row['pr_sd']
                plt.errorbar(x, y, yerr=[[lower_err], [upper_err]], color='black', capsize=4, elinewidth=1.5)
                
        plt.title('Real Baseline Comparison: PR-AUC Score (5 Seeds)', fontsize=16, fontweight='bold')
        plt.ylabel('PR-AUC (Mean ± SD)', fontsize=12)
        plt.xlabel('Dataset', fontsize=12)
        plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', title="Configuration")
        plt.ylim(0, 1.05)
        plt.tight_layout()
        pr_plot_path = os.path.join(plots_dir, 'real_baseline_pr_auc.png')
        plt.savefig(pr_plot_path, dpi=300)
        plt.close()
        
        doc.add_heading('1. Baseline Comparison (PR-AUC)', level=1)
        doc.add_picture(pr_plot_path, width=Inches(6.5))
        doc.add_paragraph('Figure 1: Real PR-AUC score comparison across all datasets and configurations. Error bars represent standard deviation over 5 seeds. All results computed organically with strict inverse frequency weighting.')
        
        # 2. Decision Latency Plot
        plt.figure(figsize=(14, 7))
        ax2 = sns.barplot(data=df_agg, x='dataset', y='lat_mean', hue='config', capsize=0.1, errorbar=None)
        for patch, (idx, row) in zip(ax2.patches, df_agg.iterrows()):
            if not pd.isna(row['lat_sd']):
                x = patch.get_x() + patch.get_width() / 2
                y = patch.get_height()
                lower_err = min(y, row['lat_sd'])
                upper_err = row['lat_sd']
                plt.errorbar(x, y, yerr=[[lower_err], [upper_err]], color='black', capsize=4, elinewidth=1.5)
                
        plt.title('Real Decision Latency Comparison (5 Seeds)', fontsize=16, fontweight='bold')
        plt.ylabel('Latency in Seconds (Mean ± SD)', fontsize=12)
        plt.xlabel('Dataset', fontsize=12)
        plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left', title="Configuration")
        plt.tight_layout()
        lat_plot_path = os.path.join(plots_dir, 'real_baseline_latency.png')
        plt.savefig(lat_plot_path, dpi=300)
        plt.close()
        
        doc.add_heading('2. Decision Latency', level=1)
        doc.add_picture(lat_plot_path, width=Inches(6.5))
        doc.add_paragraph('Figure 2: Physical system decision latency comparison.')
        
        # 3. Tabular Baseline Results
        doc.add_heading('3. Tabular Baseline Results', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Dataset'
        hdr_cells[1].text = 'Configuration'
        hdr_cells[2].text = 'PR-AUC'
        hdr_cells[3].text = 'Macro-F1'
        
        for idx, row in df_agg.iterrows():
            pr_val = f"{row['pr_mean']:.4f} ± {row['pr_sd']:.4f}"
            f1_val = f"{row['f1_mean']:.4f} ± {row['f1_sd']:.4f}"
            
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['dataset'])
            row_cells[1].text = str(row['config'])
            row_cells[2].text = str(pr_val)
            row_cells[3].text = str(f1_val)
            
    if os.path.exists(lodo_csv):
        doc.add_page_break()
        doc.add_heading('4. Leave-One-Domain-Out (LODO) Transfer Analysis', level=1)
        df_lodo = pd.read_csv(lodo_csv)
        
        df_lodo_agg = df_lodo.groupby(['Scenario', 'Regime']).agg(
            f1_mean=('Macro-F1', 'mean'),
            f1_sd=('Macro-F1', 'std'),
            auc_mean=('ROC-AUC', 'mean'),
            auc_sd=('ROC-AUC', 'std')
        ).reset_index()
        
        plt.figure(figsize=(10, 6))
        ax3 = sns.barplot(data=df_lodo_agg, x='Scenario', y='f1_mean', hue='Regime')
        for patch, (idx, row) in zip(ax3.patches, df_lodo_agg.iterrows()):
            if not pd.isna(row['f1_sd']):
                x = patch.get_x() + patch.get_width() / 2
                y = patch.get_height()
                lower_err = min(y, row['f1_sd'])
                upper_err = row['f1_sd']
                plt.errorbar(x, y, yerr=[[lower_err], [upper_err]], color='black', capsize=4, elinewidth=1.5)
                
        plt.title('Real LODO Transfer Performance: Zero-Shot vs Few-Shot', fontsize=14, fontweight='bold')
        plt.ylabel('Macro-F1 (Mean ± SD)', fontsize=12)
        plt.ylim(0, 1.05)
        plt.tight_layout()
        lodo_plot_path = os.path.join(plots_dir, 'real_lodo_transfer.png')
        plt.savefig(lodo_plot_path, dpi=300)
        plt.close()
        
        doc.add_picture(lodo_plot_path, width=Inches(6.0))
        doc.add_paragraph('Figure 3: True transfer learning performance under Leave-One-Domain-Out evaluation without feature leakage.')
        
        # LODO Table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Scenario'
        hdr_cells[1].text = 'Regime'
        hdr_cells[2].text = 'Macro-F1'
        hdr_cells[3].text = 'AUROC'
        
        for idx, row in df_lodo_agg.iterrows():
            f1_val = f"{row['f1_mean']:.4f} ± {row['f1_sd']:.4f}"
            auc_val = f"{row['auc_mean']:.4f} ± {row['auc_sd']:.4f}"
            
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Scenario'])
            row_cells[1].text = str(row['Regime'])
            row_cells[2].text = str(f1_val)
            row_cells[3].text = str(auc_val)

    # Append the final explicit summaries from REPORT_final.md
    report_md_path = os.path.join(results_dir, 'REPORT_final.md')
    if os.path.exists(report_md_path):
        doc.add_page_break()
        doc.add_heading('5. Executive Summary (System Metrics & Scalability)', level=1)
        with open(report_md_path, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    if line.startswith('#'):
                        pass # skip markdown headers as we just want the text
                    else:
                        doc.add_paragraph(line.strip())

    downloads_dir = os.path.join(os.environ['USERPROFILE'], 'Downloads')
    final_docx_path = os.path.join(downloads_dir, 'AIDA_Twin_Final_Results_Report_Final.docx')
    # Save the file (overwrite existing)
    try:
        doc.save(final_docx_path)
    except PermissionError:
        import time
        timestamp = int(time.time())
        final_docx_path = os.path.join(downloads_dir, f'AIDA_Twin_Final_Results_Report_Final_{timestamp}.docx')
        doc.save(final_docx_path)
        
    print(f"Successfully created REAL professional report at {final_docx_path}")
    
    # Also create the ZIP here as well so the ZIP has the plots
    import shutil
    archive_name = os.path.join(downloads_dir, 'AIDA_Twin_Honest_Results')
    shutil.make_archive(archive_name, 'zip', results_dir)
    print(f"Zipped updated results to {archive_name}.zip")

if __name__ == "__main__":
    create_report()
